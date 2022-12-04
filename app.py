import os
from flask import Flask, flash, request, redirect, url_for, render_template, send_from_directory, session
from werkzeug.utils import secure_filename #ファイル名保護用
from datetime import timedelta #settion管理で時間情報を用いるため

'''Docx'''
import docx  #Microsoft Word ファイルを扱うため'''
from docx.oxml import OxmlElement #Open XML ファイル形式を扱うため'''
from docx.oxml.ns import qn #XML 名前空間を扱うため'''
from docx.text.paragraph import Paragraph #Microsoft Word の文章を表すクラス'''
from docx.oxml.xmlchemy import OxmlElement
from docx import Document #Microsoft Word のドキュメントを表すクラス'''
from retry import retry #パッケージは、再試行を行うためのモジュール'''
from googletrans import Translator #Google 翻訳 API を使用して、テキストの翻訳を行うためのクラス'''
import array #配列を扱うため'''
import os #オペレーティングシステムへのアクセスを提供するモジュール'''
import sys #Python インタプリターを扱うため'''
import time #時刻を扱うため'''
import requests #Web リクエストを行うため'''
import json #JSON データを扱うため'''

''''''
UPLOAD_FOLDER = './uploads'
TRANSLATED_FOLDER = './translated'
ALLOWED_EXTENSIONS = set(['docx', 'doc'])

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['TRANSLATED_FOLDER'] = TRANSLATED_FOLDER
app.secret_key = 'user'
app.permanent_session_lifetime = timedelta(minutes=5) # -> 5分 #(days=5) -> 5日保存

#拡張子が有効かどうか確認する関数
def allowed_file(filename):
    # .があるかどうかのチェックと、拡張子の確認
    # OKなら１、だめなら0
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    return render_template("index.html")



@app.route('/upload', methods=['GET', 'POST'])

#アップロードされたファイルのURLにリダイレクトする関数
#流れ
#「/upload」URLに対してPOSTリクエストが行われた場合に実行される関数を定義する。
#POSTリクエストから「filename」という名前で送信されたファイルを取得する。
#「filename」がPOSTリクエストの中に存在するかどうかをチェックし、存在しない場合は、「POSTリクエストにファイルが含まれていません。」というエラーメッセージを表示する。
#「allowed_file」関数を使用して、アップロードされたファイルの拡張子が有効かどうかを確認する。

def upload_file():
    #リクエストメソッドがPOSTの場合
    if request.method == 'POST':
        print("POSTがリクエストされました")

        file = request.files.get('filename')

        # POSTリクエストがファイル部分を持つかどうかをチェック
        if 'filename' not in request.files:
            flash("リクエストにファイルが含まれていません。一度タブを閉じてもう一度初めからやり直してください。")

        # ユーザーがファイルを選択しない場合
        if file.filename == '':
            flash("ファイルが選択されていません。Choose a fileのボタンを押してファイルを選択してください。")

        #ファイルがあり、かつ許可されたファイル形式である場合
        if file and allowed_file(file.filename):
            securename = secure_filename(file.filename)
            print(securename)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], securename))

            #翻訳言語の選択
            lang_option = request.form.get('lang', '')
            if lang_option == "EN → JA":
                dest_lang = "ja"
            else:
                dest_lang = 'en'
            print(dest_lang)
            return redirect(url_for('translate_file', name=securename, lang=dest_lang))
        else:
            flash("ファイル形式が異なります。現在利用できるファイル形式は.docxまたは.docです。決議などをPDFでしか持っていない場合はODSからwordファイルを取得してください。")
    return render_template("index.html")


@app.route('/uploads/<lang>/translated_<name>')
def translate_file(name, lang):
    dest_lang = lang
    print(dest_lang)
    translator = Translator(service_urls=['translate.googleapis.com'])
    doc = docx.Document("./uploads/"+name)
    pg= len(doc.paragraphs) #文書内の段落数を取得
    print("処理が始まりました")

    array=[]
    for i in range(1, pg): #ループ：最初の段落であるタイトルを除いてドキュメントの各段落に実行　あとでtqdmかます
        try:
            para=doc.paragraphs[i]
        except IndexError as e:
          print(e)
          break #トライキャッチ：インデックスエラーの捕捉。パラグラフが翻訳されていればコメントが追加。もし段落が空であればその段落をスキップ

        bfori_text=para.text
        bfori_text=bfori_text.replace("\r","")
        ori_text=bfori_text.replace("\x07","")
        if str(ori_text)=="":
            z=1
            print("パラグラフ　"+str(i)+"　はテキストが存在しないのでパスしました。")
        else:
            z=0
            translation = translator.translate(ori_text, dest=dest_lang)
            array.extend([ori_text, translation.text])
            time.sleep(0.0005)
            print("Google Translate :: "+ori_text)
        if z==0:
            try:
                run = para.add_run(' ')
                run.add_comment(translation.text,author='DR_translator',initials= 'KT')
            except:
                print("何らかの理由により、パラグラフ　"+str(i)+"　はコメントを追加することができませんでした。")

    core_properties = doc.core_properties
    lastpara = doc.add_paragraph(" ")
    lastrun = lastpara.add_run(' ')
    lastrun.add_comment("Document_Revision: "+str(core_properties.revision)+"\n"+"このファイルはDR_translatorによって自動翻訳されています。訳文のコメントを削除したい場合は、校閲→コメント→ドキュメント内のすべてのコメントを削除を押下してください。",author='DR_translator',initials= 'KT')
    downloadfile = "translated_"+name
    doc.save(os.path.join(app.config['TRANSLATED_FOLDER'], downloadfile))
    return send_from_directory(app.config["TRANSLATED_FOLDER"], downloadfile)

if __name__ == '__main__':
    app.debug = True
    app.run(host='localhost')
